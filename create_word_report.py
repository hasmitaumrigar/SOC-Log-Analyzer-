from docx import Document
from docx.shared import Inches
import os

# Create a new Word document
doc = Document()
doc.add_heading("SOC Log Analyzer Project", 0)

# Introduction
doc.add_heading("Overview", level=1)
doc.add_paragraph(
    "This project is a Security Operations Center (SOC) Log Analyzer built in Python. "
    "It parses Apache web server logs, detects security incidents, visualizes trends, "
    "and generates a professional PDF report. This tool helps SOC analysts monitor "
    "threats like brute-force attacks, DoS attempts, and error spikes."
)

# Features
doc.add_heading("Features", level=1)
features = [
    "Log Parsing: Extracts IP, Timestamp, Method, URL, Status, Size, and Minute.",
    "Brute-force Detection: Flags IPs with ≥3 failed login attempts.",
    "DoS Detection: Flags IPs with ≥5 requests per minute.",
    "Error Spike Detection: Flags IPs generating ≥3 HTTP 404/500 errors.",
    "Visualization: Top attacking IPs, failed login attempts, and error spikes charts.",
    "PDF Report: Generates a 2-page professional report with alerts and dashboard."
]
for f in features:
    doc.add_paragraph(f"• {f}")

# Technologies
doc.add_heading("Technologies Used", level=1)
techs = ["Python 3", "pandas", "matplotlib", "fpdf"]
for t in techs:
    doc.add_paragraph(f"• {t}")

# Project Structure
doc.add_heading("Project Structure", level=1)
structure = [
    "analyzer.py – Main script",
    "detection.py – Functions to detect brute-force, DoS, error spikes",
    "visualization.py – Charts and dashboard plotting",
    "report.py – PDF generation",
    "logs/apache.log – Sample log file",
    "output/ – Output folder containing alerts.csv, dashboard.png, SOC_Report.pdf"
]
for s in structure:
    doc.add_paragraph(f"• {s}")

# Screenshots
doc.add_heading("Screenshots", level=1)
screenshots_folder = "env"  # replace with your folder path
if os.path.exists(screenshots_folder):
    for file in os.listdir(screenshots_folder):
        if file.lower().endswith(('.png', '.jpg', '.jpeg')):
            doc.add_paragraph(file)
            doc.add_picture(os.path.join(screenshots_folder, file), width=Inches(5))
else:
    doc.add_paragraph("No screenshots folder found.")

# Learning / Mistake
doc.add_heading("Learning / Mistake", level=1)
doc.add_paragraph(
    "Initially, I added the full datetime as the 'Minute' column in logs, "
    "which made the PDF table unreadable. I corrected it by formatting it as HH:MM. "
    "This taught me the importance of data formatting for reporting and visualization."
)

# Conclusion
doc.add_heading("Conclusion", level=1)
doc.add_paragraph(
    "This project demonstrates end-to-end log analysis, alert detection, "
    "visualization, and reporting without external devices. "
    "It is suitable for internship submissions and portfolio demonstrations."
)

# Save the Word file
output_file = "SOC_Log_Analyzer_Report.docx"
doc.save(output_file)
print(f"Word report generated: {output_file}")